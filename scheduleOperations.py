from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.opc.exceptions import PackageNotFoundError
import os, time

ADD_SCHEDS = []


def insert_sched(
    day: str, subject: str, start_time: str, end_time: str, document
) -> int:
    table = document.tables[0]
    header_cells = table.rows[0].cells

    for column in range(len(header_cells)):
        # look for a certain column of a certain day
        if header_cells[column].text == day:
            for row in range(1, len(table.columns[column].cells)):
                # check if the cell is empty
                if table.columns[column].cells[row].text.strip() == "":
                    # put text in that cell and center it in the cell
                    table.columns[column].cells[
                        row
                    ].text = f"{subject}\n({start_time} - {end_time})"
                    p = table.columns[column].cells[row].paragraphs[0]
                    p.alignment = WD_TABLE_ALIGNMENT.CENTER
                    return 0

                # if it's the last cell in the column and it is not blank, then create a new row and place the text there
                if row == len(table.columns[column].cells) - 1:
                    table.add_row().cells
                    table.columns[column].cells[
                        row + 1
                    ].text = f"{subject}\n({start_time} - {end_time})"
                    p = table.columns[column].cells[row + 1].paragraphs[0]
                    p.alignment = WD_TABLE_ALIGNMENT.CENTER
                    return 0
    return 1


def sortSched(schedList: list) -> int:
    if len(schedList) == 0:
        return 1

    schedList.sort(key=lambda x: x[2])
    for i in range(len(schedList)):
        schedList[i][2] = time.strftime("%I:%M %p", schedList[i][2])

    return 0


def generate_sched(saveFileName: str, loadFileName: str) -> int:
    sortFailed = sortSched(ADD_SCHEDS)
    if sortFailed:
        return 4

    try:
        # open the template file
        # make sure not to overwrite template.docx!!
        document = Document(f"{os.getcwd()}/Schedules/{loadFileName}")

        for day, subject, startTime, endTime in ADD_SCHEDS:
            insertFailed = insert_sched(day, subject, startTime, endTime, document)

        # if failed to insert sched for sum reason
        if insertFailed:
            return 1

        document.save(f"Schedules/{saveFileName}.docx")

    except PackageNotFoundError:
        print("Invalid document.")
        return 1

    except PermissionError:
        for i in range(len(ADD_SCHEDS)):
            # revert the start time in string format back to struct_time format since sortSched() already converted the start time to string format
            # this is to avoid a type error in converting to and from struct_time format
            ADD_SCHEDS[i][2] = time.strptime(ADD_SCHEDS[i][2], "%I:%M %p")
        return 2

    # clear the ADD_SCHEDS array
    ADD_SCHEDS.clear()

    return 0


if __name__ == "__main__":
    # ignore
    pass
